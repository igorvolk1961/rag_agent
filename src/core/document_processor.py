"""
Document processing functionality using LangChain
"""

import logging
import time
from pathlib import Path
from typing import List, Dict, Any, Optional

from langchain_community.document_loaders import TextLoader, PyPDFLoader
from docx import Document as DocxDocument
from langchain_core.documents import Document as LangChainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter, TextSplitter

from models.schemas import Document, Chunk
from modules.parser.smart_chunker_adapter import SmartChunkerAdapter


class DocumentProcessor:
    """Document processor using LangChain loaders and splitters"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200,
                 use_smart_chunker: bool = False, smart_chunker_config: Optional[str] = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_smart_chunker = use_smart_chunker
        self.logger = logging.getLogger(__name__)

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

        # Initialize SmartChunker if enabled
        self.smart_chunker = None
        if use_smart_chunker:
            try:
                self.smart_chunker = SmartChunkerAdapter(smart_chunker_config)
                self.logger.info("SmartChunker инициализирован успешно")
            except Exception as e:
                self.logger.warning(f"Не удалось инициализировать SmartChunker: {e}")
                self.logger.warning("Будет использован стандартный чанкер")
                self.use_smart_chunker = False

    def load_document(self, file_path: Path) -> LangChainDocument:
        """Load document using LangChain loaders"""
        try:
            suffix = file_path.suffix.lower()

            if suffix == '.txt':
                loader = TextLoader(str(file_path), encoding='utf-8')
                documents = loader.load()
            elif suffix == '.pdf':
                loader = PyPDFLoader(str(file_path))
                documents = loader.load()
            elif suffix == '.docx':
                documents = self._load_docx_with_python_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")

            if not documents:
                raise ValueError(f"No content found in {file_path}")

            # Combine multiple pages into single document
            if len(documents) > 1:
                combined_content = "\n\n".join([doc.page_content for doc in documents])
                combined_metadata = documents[0].metadata.copy()
                combined_metadata.update({
                    "total_pages": len(documents),
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size
                })
                document = LangChainDocument(
                    page_content=combined_content,
                    metadata=combined_metadata
                )
            else:
                document = documents[0]
                document.metadata.update({
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size
                })

            self.logger.info(f"Loaded document: {file_path}")
            return document

        except Exception as e:
            self.logger.error(f"Error loading document {file_path}: {e}")
            raise

    def _load_docx_with_python_docx(self, file_path: Path) -> List[LangChainDocument]:
        """Load DOCX document using python-docx for better table and structure handling"""
        try:
            # Load document with python-docx
            doc = DocxDocument(str(file_path))

            # Extract text content paragraph by paragraph
            text_content = []
            tables_content = []

            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Process tables
            for i, table in enumerate(doc.tables):
                table_text = f"\n\n--- ТАБЛИЦА {i+1} ---\n"
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        table_text += " | ".join(row_cells) + "\n"
                tables_content.append(table_text)

            # Combine all content
            full_content = "\n".join(text_content)
            if tables_content:
                full_content += "\n" + "\n".join(tables_content)

            # Create LangChain document
            document = LangChainDocument(
                page_content=full_content,
                metadata={
                    "source": str(file_path),
                    "file_type": "docx",
                    "tables_count": len(tables_content),
                    "paragraphs_count": len([p for p in doc.paragraphs if p.text.strip()])
                }
            )

            return [document]

        except Exception as e:
            self.logger.error(f"Error loading DOCX with python-docx {file_path}: {e}")
            raise

    def process_document(self, file_path: Path) -> Document:
        """Process a document and return Document object"""
        try:
            # Load document using LangChain
            langchain_doc = self.load_document(file_path)

            # Convert to our Document format
            document = Document(
                id=file_path.stem,
                title=file_path.stem,
                content=langchain_doc.page_content,
                metadata={
                    "type": file_path.suffix.lower().lstrip('.'),
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size,
                    **langchain_doc.metadata
                }
            )

            self.logger.info(f"Processed document: {file_path}")
            return document

        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {e}")
            raise

    def chunk_document(self, document: Document) -> List[Chunk]:
        """Split document into chunks using LangChain text splitter"""
        try:
            # Create LangChain document for splitting
            langchain_doc = LangChainDocument(
                page_content=document.content,
                metadata=document.metadata
            )

            # Split document
            chunks = self.text_splitter.split_documents([langchain_doc])

            # Convert to our Chunk format
            result_chunks = []
            for i, chunk in enumerate(chunks):
                chunk_obj = Chunk(
                    id=f"{document.id}_chunk_{i}",
                    document_id=document.id,
                    content=chunk.page_content,
                    start_pos=0,  # LangChain doesn't track positions
                    end_pos=len(chunk.page_content),
                    metadata={
                        "chunk_id": i,
                        "document_title": document.title,
                        "document_type": document.metadata.get("type", "unknown"),
                        **chunk.metadata
                    }
                )
                result_chunks.append(chunk_obj)

            self.logger.info(f"Created {len(result_chunks)} chunks for document {document.id}")
            return result_chunks

        except Exception as e:
            self.logger.error(f"Error chunking document {document.id}: {e}")
            raise

    def process_and_chunk(self, file_path: Path) -> List[Chunk]:
        """Process document and return chunks in one step"""
        document = self.process_document(file_path)
        return self.chunk_document(document)

    def process_multiple_documents(self, file_paths: List[Path]) -> List[Chunk]:
        """Process multiple documents and return all chunks"""
        start_time = time.time()
        all_chunks = []

        for i, file_path in enumerate(file_paths, 1):
            try:
                file_start = time.time()
                chunks = self.process_and_chunk(file_path)
                file_time = time.time() - file_start
                all_chunks.extend(chunks)
                self.logger.info(f"📄 Документ {i}/{len(file_paths)}: {file_path.name} - {len(chunks)} чанков за {file_time:.2f}с")
            except Exception as e:
                file_time = time.time() - file_start
                self.logger.error(f"Error processing {file_path} after {file_time:.2f}с: {e}")
                continue

        total_time = time.time() - start_time
        self.logger.info(f"✅ Обработано {len(file_paths)} документов, создано {len(all_chunks)} чанков за {total_time:.2f}с")
        return all_chunks

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        if self.use_smart_chunker and self.smart_chunker:
            return self.smart_chunker.get_supported_extensions()
        return ['.txt', '.pdf', '.docx']

    def process_document_with_smart_chunker(self, file_path: Path) -> Document:
        """
        Process document using SmartChunker

        Args:
            file_path: Path to document

        Returns:
            Document object with hierarchical chunks
        """
        if not self.use_smart_chunker or not self.smart_chunker:
            raise ValueError("SmartChunker не инициализирован")

        return self.smart_chunker.process_document(str(file_path))

    def process_multiple_documents_with_smart_chunker(self, file_paths: List[Path]) -> List[Document]:
        """
        Process multiple documents using SmartChunker

        Args:
            file_paths: List of document paths

        Returns:
            List of Document objects
        """
        if not self.use_smart_chunker or not self.smart_chunker:
            raise ValueError("SmartChunker не инициализирован")

        start_time = time.time()
        self.logger.info(f"🔄 SmartChunker обрабатывает {len(file_paths)} документов...")

        documents = self.smart_chunker.process_documents([str(fp) for fp in file_paths])

        total_time = time.time() - start_time
        total_chunks = sum(len(doc.chunks) for doc in documents)
        self.logger.info(f"✅ SmartChunker обработал {len(file_paths)} документов, создал {total_chunks} чанков за {total_time:.2f}с")

        return documents

    def is_smart_chunker_available(self) -> bool:
        """Check if SmartChunker is available"""
        return self.use_smart_chunker and self.smart_chunker is not None

    def is_supported_file(self, file_path: Path) -> bool:
        """Check if file type is supported"""
        return file_path.suffix.lower() in self.get_supported_extensions()

    def set_text_splitter(self, text_splitter: TextSplitter) -> None:
        """Set custom text splitter"""
        self.text_splitter = text_splitter
        self.logger.info("Updated text splitter")

    def get_text_splitter(self) -> TextSplitter:
        """Get current text splitter"""
        return self.text_splitter
